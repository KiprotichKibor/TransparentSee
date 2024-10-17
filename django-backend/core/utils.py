from io import BytesIO
import re

from weasyprint import HTML
from .models import Investigation, Contribution
from django.template.loader import render_to_string
from django.utils import timezone
from docx import Document

def enhanced_content_moderation(content):
    '''
    Basic moderation function to check for offensive words in content
    '''
    offensive_words = [
        'kuma',
        'matako',
        'kumamako',
        'mafi',
        'mafisi',
        'mavi',
        'dinywa',
        'jidinye',
        'pussy'
    ]
    
    for word in offensive_words:
        if re.search(r'\b' + word + r'\b', content, re.IGNORECASE):
            return False
    return True
    
def generate_report_content(case_report, template='default'):
    '''
    Generate a report content from a case report
    '''
    context = {
        'case_report': case_report,
        'investigation': case_report.investigation,
        'report': case_report.investigation.report,
        'contributions': case_report.investigation.contributions.all(),
    }

    template_name = f'core/case_report_{template}.html'
    return render_to_string(template_name, context)

def generate_report_pdf(case_report):
    '''
    Generate a PDF file from a case report
    '''
    html_content = generate_report_content(case_report)
    html = HTML(string=html_content)
    return html.write_pdf()

def generate_report_docx(case_report):
    '''
    Generate a DOCX file from a case report
    '''
    document = Document()
    document.add_heading(f"Case Report: {case_report.investigation.report.title}", 0)

    document.add_heading('Report Details', level=1)
    document.add_paragraph(f"Description: {case_report.investigation.report.description}")
    document.add_paragraph(f"Region: {case_report.investigation.report.region.name}")
    document.add_paragraph(f"Date Reported: {case_report.investigation.report.created_at.strftime('%B %d, %Y')}")

    document.add_heading('Investigation Details', level=1)
    document.add_paragraph(f"Status: {case_report.investigation.get_status_display()}")
    document.add_paragraph(f"Date Started: {case_report.investigation.created_at.strftime('%B %d, %Y')}")

    document.add_heading('Contributions', level=1)
    for contribution in case_report.investigation.contributions.all():
        document.add_paragraph(f"Type: {contribution.get_contribution_type_display()}")
        document.add_paragraph(f"Content: {contribution.content}")
        document.add_paragraph(f"Created At: {contribution.created_at.strftime('%B %d, %Y')}")
        document.add_paragraph('____________________')

    bio = BytesIO()
    document.save(bio)
    return bio.getvalue()